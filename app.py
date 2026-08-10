"""
app.py — FastAPI backend for the PDF RAG Assistant web UI.

Endpoints:
  POST /upload    — receive PDF, run ingestion in background thread
  POST /ask       — answer a question using the RAG agent
  GET  /status    — ingestion progress / readiness check
  GET  /          — serves static/index.html (the web UI)

Run with:
    uvicorn app:app --reload --port 8000
Then open http://localhost:8000
"""

import os
import shutil
import asyncio
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, TextLoader
from langchain_core.documents import Document
import docx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from config import vectorstore, index as pinecone_index, add_documents_safely
from rag_agent import ask, ask_with_sources

# ── App ─────────────────────────────────────────────────────────────────────
app = FastAPI(title="Document RAG Assistant")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

# Thread pool: one thread for ingest, one for LLM inference
executor = ThreadPoolExecutor(max_workers=2)

# Shared single-user state
state: dict = {
    "status": "idle",      # idle | processing | ready | error
    "filename": None,
    "pages": 0,
    "chunks": 0,
    "error": None,
}


# ── Ingestion (blocking — runs in thread pool) ───────────────────────────────
def _run_ingest(file_path: str) -> None:
    try:
        path_lower = file_path.lower()
        pages = []
        
        # 1. Structure-aware Document Loading
        if path_lower.endswith(".docx") or path_lower.endswith(".doc"):
            try:
                doc = docx.Document(file_path)
                current_page_text = []
                page_num = 1
                for p in doc.paragraphs:
                    text = p.text.strip()
                    if not text:
                        continue
                    current_page_text.append(text)
                    combined_len = sum(len(t) for t in current_page_text)
                    if combined_len >= 2500 or (text.isupper() and len(text) > 4):
                        pages.append(Document(
                            page_content="\n\n".join(current_page_text),
                            metadata={"source": file_path, "page": page_num}
                        ))
                        current_page_text = []
                        page_num += 1
                if current_page_text:
                    pages.append(Document(
                        page_content="\n\n".join(current_page_text),
                        metadata={"source": file_path, "page": page_num}
                    ))
            except Exception:
                loader = Docx2txtLoader(file_path)
                pages = loader.load()
        elif path_lower.endswith(".txt"):
            loader = TextLoader(file_path, encoding="utf-8")
            raw_docs = loader.load()
            pages = []
            for doc_idx, raw_doc in enumerate(raw_docs, 1):
                paragraphs = raw_doc.page_content.split("\n\n")
                curr_txt = []
                p_num = 1
                for para in paragraphs:
                    curr_txt.append(para)
                    if sum(len(p) for p in curr_txt) >= 2500:
                        pages.append(Document(
                            page_content="\n\n".join(curr_txt),
                            metadata={"source": file_path, "page": p_num}
                        ))
                        curr_txt = []
                        p_num += 1
                if curr_txt:
                    pages.append(Document(
                        page_content="\n\n".join(curr_txt),
                        metadata={"source": file_path, "page": p_num}
                    ))
        else:
            loader = PyPDFLoader(file_path)
            pages = loader.load()

        state["pages"] = len(pages)

        # 2. Structure-aware Chunking Strategy (~900 tokens = 2700 chars, ~120 tokens overlap = 360 chars)
        CHUNK_SIZE = 2700     # Approx 900 tokens
        CHUNK_OVERLAP = 360   # Approx 120 tokens

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            separators=[
                "\n# ",           # Markdown H1 / Major Section Title
                "\n## ",          # Markdown H2 / Subsection Title
                "\n### ",         # Markdown H3 / Sub-subsection Title
                "\n\n\n",         # Multi-paragraph logical break
                "\n\n",           # Paragraph boundary
                "\n",             # Line break
                " ",              # Space
                ""
            ],
        )
        chunks = splitter.split_documents(pages)

        # 3. Enrich chunk metadata with section titles & chunk indices
        for idx, chunk in enumerate(chunks, 1):
            chunk.metadata["chunk_index"] = idx
            lines = [l.strip() for l in chunk.page_content.split("\n") if l.strip()]
            for line in lines[:4]:
                if (line.startswith(("#", "SECTION", "Section", "PART", "Part")) or
                    (len(line) < 80 and line.isupper() and len(line) > 4) or
                    (len(line) > 3 and line[:3].strip().replace(".", "").isdigit())):
                    chunk.metadata["section"] = line[:100]
                    break

        state["chunks"] = len(chunks)

        # 4. Ingestion Diagnostics & Logging
        avg_chars = sum(len(c.page_content) for c in chunks) // max(1, len(chunks))
        avg_tokens = int(avg_chars / 3.5)
        print("=" * 60)
        print(f"[INGESTION DIAGNOSTICS] File: {Path(file_path).name}")
        print(f" -> Total Pages Processed: {len(pages)}")
        print(f" -> Total Chunks Created: {len(chunks)}")
        print(f" -> Avg Chunk Size: {avg_chars} chars (~{avg_tokens} tokens)")
        print(f" -> Config: chunk_size={CHUNK_SIZE} chars (~900 tokens), overlap={CHUNK_OVERLAP} chars (~120 tokens)")
        print(f" -> Vector DB: Pinecone FastEmbed ONNX (384-dim)")
        print("=" * 60)

        # Clear previous vectors so re-uploading doesn't create duplicates
        try:
            pinecone_index.delete(delete_all=True)
        except Exception:
            pass  # index may be empty — that's fine

        # Embed + upsert (Instant FastEmbed local embeddings with zero rate limits)
        add_documents_safely(vectorstore, chunks, batch_size=100, delay_seconds=0.0)
        state["status"] = "ready"

    except Exception as exc:
        state["status"] = "error"
        state["error"]  = str(exc)


# ── Routes — defined BEFORE the static mount ────────────────────────────────
@app.post("/upload")
async def upload(file: UploadFile = File(...)):
    allowed_exts = (".pdf", ".docx", ".doc", ".txt")
    if not file.filename.lower().endswith(allowed_exts):
        raise HTTPException(status_code=400, detail="Only PDF, Word (.docx/.doc), and Text (.txt) files are supported.")

    dest = UPLOAD_DIR / file.filename
    with open(dest, "wb") as f:
        shutil.copyfileobj(file.file, f)

    state.update({
        "status":   "processing",
        "filename": file.filename,
        "pages":    0,
        "chunks":   0,
        "error":    None,
    })

    loop = asyncio.get_running_loop()
    loop.run_in_executor(executor, _run_ingest, str(dest))

    return {"message": "Upload received. Ingestion started.", "filename": file.filename}


class QuestionBody(BaseModel):
    question: str


@app.post("/ask")
async def ask_question(body: QuestionBody):
    if state["status"] != "ready":
        raise HTTPException(
            status_code=400,
            detail="Document not ready yet. Upload a document and wait for processing to complete.",
        )
    if not body.question.strip():
        raise HTTPException(status_code=400, detail="Question cannot be empty.")

    try:
        loop = asyncio.get_running_loop()
        res = await loop.run_in_executor(executor, ask_with_sources, body.question.strip())
        return {
            "answer": res["answer"],
            "filename": state["filename"],
            "sources": res.get("sources", [])
        }
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))


@app.get("/status")
async def get_status():
    return state


# ── Static / Frontend UI — must be mounted LAST ──────────────────────────────
frontend_dir = "frontend" if Path("frontend").exists() else "static"
app.mount("/", StaticFiles(directory=frontend_dir, html=True), name="static")

